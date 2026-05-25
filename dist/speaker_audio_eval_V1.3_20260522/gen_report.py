# -*- coding: utf-8 -*-
"""根据结构化结果生成 Word 声学评测报告。"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from markdown_report import DIMENSION_KEYS, append_section_six_docx


def generate_report(result_data: dict, save_path: str = "音效对比测试报告.docx") -> str:
    """
    result_data 字段：
      test_name, test_device, ref_device, 总体平均分, 综合结论, 单行评述（可选）,
      comparison_mode, cross_session,
      明细(list[dict])：每项含 节目、分组、五维键、对比总结/综合分、专业点评
    """
    try:
        doc = Document()

        title = doc.add_heading("终端设备喇叭音效 AI 辅助评测报告", 0)
        title.alignment = 1

        doc.add_paragraph(f"测试名称：{result_data.get('test_name', '')}")
        _em = str(result_data.get("eval_model") or "").strip()
        if _em:
            doc.add_paragraph(f"评测大模型：{_em}")
        doc.add_paragraph(f"被测设备：{result_data.get('test_device', '')}")
        doc.add_paragraph(f"对比/参考：{result_data.get('ref_device', '')}")
        n_tracks = result_data.get("音源条数") or str(len(result_data.get("明细") or []))
        doc.add_paragraph(f"测试音源：共 {n_tracks} 条（由 assets/test_audio 自动扫描排序）")
        if result_data.get("comparison_mode"):
            if result_data.get("cross_session"):
                doc.add_paragraph(
                    "评分说明：**跨会话**刺激比较 — 两次测试各连一台机、单路麦克分时录制后按音源对齐；"
                    "评分规则与同一会话内双机顺序录制**相同**：五维仅整数 -3～+3（正=被测优，负=被测劣，0=相当），"
                    "禁止小数与区间。（两次会话间环境应尽量一致。）"
                )
            else:
                doc.add_paragraph(
                    "评分说明：多设备 **刺激比较** — 对比【被测设备】与【对比设备】在同一音源下内置喇叭经麦克风采集的主观听感；"
                    "五维分差取值仅为整数 -3～+3（正=被测优，负=被测劣，0=相当），禁止小数与区间。"
                    "（录音存在环境与摆放偏差。）"
                )
        else:
            doc.add_paragraph(
                "评分说明：各维度为模型给出的 1–10 分主观评价（录音经麦克风采集，存在环境偏差）。"
            )
        doc.add_paragraph()

        avg_label = result_data.get("平均分说明") or "总体平均分"
        doc.add_paragraph(f"{avg_label}：{result_data.get('总体平均分', 'N/A')}")
        doc.add_paragraph(f"综合结论：{result_data.get('综合结论', '')}")
        if result_data.get("单行评述"):
            doc.add_paragraph(f"整体评述：{result_data.get('单行评述', '')}")
        doc.add_paragraph()

        dim_avgs = result_data.get("维度平均分") or {}
        if isinstance(dim_avgs, dict) and dim_avgs:
            doc.add_paragraph("维度平均分（全节目）：")
            for k in DIMENSION_KEYS:
                v = dim_avgs.get(k)
                if v is not None:
                    doc.add_paragraph(f"  - {k}：{v}")
            if result_data.get("总平均分") is not None:
                doc.add_paragraph(f"  - 所有维度总平均分：{result_data.get('总平均分')}")
            doc.add_paragraph()

        cmp_mode = bool(result_data.get("comparison_mode"))
        if cmp_mode:
            keys = ["节目", "分组", *DIMENSION_KEYS, "对比总结", "专业点评"]
        else:
            keys = ["节目", "分组", *DIMENSION_KEYS, "综合分", "专业点评"]

        table = doc.add_table(rows=1, cols=len(keys))
        hdr_cells = table.rows[0].cells
        for i, k in enumerate(keys):
            hdr_cells[i].text = k

        for item in result_data.get("明细") or []:
            row_cells = table.add_row().cells
            for i, k in enumerate(keys):
                v = item.get(k)
                if v is None:
                    row_cells[i].text = "—"
                elif isinstance(v, float) and cmp_mode and abs(v - round(v)) < 1e-9:
                    row_cells[i].text = str(int(round(v)))
                elif isinstance(v, float):
                    row_cells[i].text = f"{v:.2f}"
                else:
                    row_cells[i].text = str(v)

        append_section_six_docx(
            doc,
            comparison_mode=cmp_mode,
            dim_avgs=dim_avgs if isinstance(dim_avgs, dict) else {},
            grand=result_data.get("总平均分"),
            rows=list(result_data.get("明细") or []),
        )

        out = Path(save_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return str(out)
    except Exception as exc:
        raise RuntimeError(f"生成 Word 报告失败: {exc}") from exc
