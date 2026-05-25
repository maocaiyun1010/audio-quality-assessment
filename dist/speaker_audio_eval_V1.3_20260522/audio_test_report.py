# -*- coding: utf-8 -*-
"""
《喇叭测试报告》Word 自动生成：从流水线 analysis JSON 或规范 dict 生成六章 Word。

主入口：
  - ``build_report_payload_from_analysis``：analysis JSON → 报告数据 dict
  - ``generate_audio_test_report``：报告数据 dict → ``音效对比报告_YYYYMMDD.docx``
"""
from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Mapping, Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from eval_source_summary import build_per_track_rows
from markdown_report import (
    DIMENSION_KEYS,
    _DIM_TUNING_DUT_WORSE,
    _DIM_TUNING_REF_WORSE,
    _pick_dims_sorted,
    _track_dim_sum,
    build_dim_conclusion_notes,
    compute_dimension_statistics,
    grand_conclusion_note,
    one_line_summary_comparison,
)

# ---------------------------------------------------------------------------
# 报告说明（第六章固定文案）
# ---------------------------------------------------------------------------
_REPORT_DISCLAIMER_PARAS: tuple[str, str] = (
    "本报告由音频自动化评测流水线根据 Dify 大模型对麦克风采集录音的主观听感分析结果自动生成，"
    "五维分差为刺激比较标尺（−3～+3，正值表示测试机优于对比机）。数据来源于当次会话 "
    "output/analysis 下 analysis JSON，节目级明细与维度聚合均与 Web UI 展示一致。",
    "报告中的调音与优化建议仅作工程参考，实际调音须结合硬件结构、腔体设计与产线一致性要求，"
    "并在标准音量、大音量两档对典型优势/短板节目复听验证后再实施。",
)

_REF_OPTIMIZATION_DEFAULT: tuple[str, str] = (
    "在测试机已占优的维度上，对比机可参照测试机约 2–5kHz 人声齿音窗与约 60–150Hz "
    "动态余量取向做小幅校准，避免整体抬升音量引入新的失真。",
    "保持与测试机相同的摆位、环境与采集链路，对 TOP 短板节目各复测一轮，"
    "确认分差趋势可重复、非偶然摆位或环境差异导致。",
)


# ---------------------------------------------------------------------------
# 从 analysis JSON 构建报告 payload
# ---------------------------------------------------------------------------
def _safe_float(v: Any) -> float | None:
    if v is None or v == "" or v == "—":
        return None
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


def _program_short_name(row: Mapping[str, Any]) -> str:
    name = str(row.get("音源名称") or row.get("节目") or "").strip()
    if name and name != "—":
        return name
    stim = str(row.get("stimulus") or "").strip()
    if stim:
        from eval_source_summary import display_source_name_from_stimulus

        return display_source_name_from_stimulus(stim) or stim
    return "无数据"


def _snippet_from_row(row: Mapping[str, Any], max_len: int = 60) -> str:
    for key in ("对比总结", "综合评价", "专业点评"):
        t = str(row.get(key) or "").strip()
        if t and t != "—":
            t = re.sub(r"\s+", "", t)
            return t[:max_len] + ("…" if len(t) > max_len else "")
    return "无数据"


def build_report_payload_from_analysis(
    analysis: Mapping[str, Any],
    *,
    test_device: str = "",
    ref_device: str = "",
    test_name: str = "",
) -> dict[str, Any]:
    """
    将流水线 ``analysis`` JSON（含 ``tracks`` / ``devices`` 等）转为
    ``generate_audio_test_report`` 所需结构；缺失字段自动推算。
    """
    if analysis.get("total_score_avg") is not None and analysis.get("dimension_scores"):
        payload = dict(analysis)
        payload.setdefault("report_date", datetime.now().strftime("%Y年%m月%d日"))
        return payload

    rows = build_per_track_rows(analysis)
    if not rows:
        for item in analysis.get("tracks") or []:
            if not item.get("ok"):
                continue
            p = item.get("parsed") if isinstance(item.get("parsed"), dict) else {}
            if not any(k in p for k in DIMENSION_KEYS):
                continue
            r: dict[str, Any] = {"音源名称": _program_short_name(item)}
            for k in DIMENSION_KEYS:
                r[k] = _safe_float(p.get(k)) or 0.0
            for key in ("对比总结", "综合评价", "专业点评", "综合结论"):
                if p.get(key):
                    r[key] = p.get(key)
            rows.append(r)

    dim_avgs, grand = compute_dimension_statistics(rows)
    dimension_scores: dict[str, float] = {}
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is not None:
            dimension_scores[k] = round(float(a), 2)

    total = analysis.get("total_score_avg")
    if total is None:
        total = grand
    if total is not None:
        total = round(float(total), 2)

    devices = list(analysis.get("devices") or [])
    dut = test_device.strip()
    ref = ref_device.strip()
    if not dut and len(devices) >= 1:
        dut = str(devices[0].get("label") or devices[0].get("slot") or "被测设备")
    if not ref and len(devices) >= 2:
        ref = str(devices[1].get("label") or devices[1].get("slot") or "对比设备")
    if not dut:
        dut = "被测设备"
    if not ref:
        ref = "对比设备"

    cross = bool(analysis.get("cross_session"))
    mode = (
        "跨会话刺激比较法（−3～+3 分差，正值=测试机优于对比机）"
        if cross
        else "刺激比较法（−3～+3 分差，正值=测试机优于对比机）"
    )

    ranked_adv = sorted(rows, key=_track_dim_sum, reverse=True)
    ranked_weak = sorted(rows, key=_track_dim_sum)
    advantage_programs: list[dict[str, Any]] = []
    weakness_programs: list[dict[str, Any]] = []
    for r in ranked_adv[:3]:
        s = _track_dim_sum(r)
        if s <= 0.05:
            continue
        advantage_programs.append(
            {
                "name": _program_short_name(r),
                "total_diff": round(s, 2),
                "highlight": _snippet_from_row(r),
            }
        )
    for r in ranked_weak[:3]:
        s = _track_dim_sum(r)
        if s >= -0.05:
            continue
        weakness_programs.append(
            {
                "name": _program_short_name(r),
                "total_diff": round(s, 2),
                "highlight": _snippet_from_row(r),
            }
        )
    if not advantage_programs:
        advantage_programs = [{"name": "无数据", "total_diff": 0.0, "highlight": "无数据"}]
    if not weakness_programs:
        weakness_programs = [{"name": "无数据", "total_diff": 0.0, "highlight": "无数据"}]

    neg_dims = _pick_dims_sorted(dim_avgs, positive=False, limit=5)
    dut_priorities: list[dict[str, Any]] = []
    for i, (k, v) in enumerate(
        sorted(neg_dims, key=lambda t: abs(t[1]), reverse=True)[:5],
        start=1,
    ):
        dut_priorities.append(
            {
                "priority": i,
                "dimension": k,
                "score": round(v, 2),
                "suggestion": _DIM_TUNING_DUT_WORSE.get(k, f"针对{k}维度优化。"),
            }
        )
    if not dut_priorities:
        dut_priorities = [
            {
                "priority": 1,
                "dimension": "综合",
                "score": 0.0,
                "suggestion": "五维分差整体接近，建议在典型节目上保持现有取向并做小范围 EQ 复核。",
            }
        ]

    pos_dims = _pick_dims_sorted(dim_avgs, positive=True, limit=3)
    highlights: list[str] = []
    for k, v in pos_dims:
        short = {"声音响度": "响度", "人声清晰度": "人声", "听感舒适度": "舒适度"}.get(k, k)
        highlights.append(f"{k}维度平均分差 {v:+.2f}，测试机在{short}方面相对领先。")
    tips: list[str] = []
    seen: set[str] = set()
    for r in rows:
        for key in ("对比总结", "综合评价", "专业点评"):
            t = str(r.get(key) or "").strip()
            if len(t) >= 12 and t not in seen and t != "—":
                seen.add(t)
                tips.append(t[:100] + ("…" if len(t) > 100 else ""))
                break
        if len(tips) >= 2:
            break
    highlights = (highlights + tips)[:3]
    if not highlights:
        highlights = ["无数据"]

    weaknesses: list[str] = []
    dim_notes = build_dim_conclusion_notes(dim_avgs, comparison_mode=True)
    for k, v in neg_dims[:3]:
        note = dim_notes.get(k, "表现偏弱")
        weaknesses.append(
            f"{k}（平均分差 {v:+.2f}）：{note}；建议优先对照短板节目复听验证。"
        )
    if not weaknesses:
        weaknesses = ["无显著短板维度，整体与对比机接近。"]

    from markdown_report import render_pairwise_comprehensive_evaluation_md

    summary = str(analysis.get("summary") or "").strip()
    if not summary:
        summary = render_pairwise_comprehensive_evaluation_md(dim_avgs, grand, rows)

    retest: list[str] = []
    if advantage_programs and advantage_programs[0].get("name") != "无数据":
        retest.append(
            f"优势场景复测：{advantage_programs[0]['name']}（确认领先幅度可重复）"
        )
    if weakness_programs and weakness_programs[0].get("name") != "无数据":
        retest.append(
            f"短板场景复测：{weakness_programs[0]['name']}（标准音量 + 大音量两档）"
        )
    if len(retest) < 2:
        retest.append("多曲目连续播放场景下的音量一致性与齿音刺激度复核。")
    retest = retest[:2]

    grand_note = grand_conclusion_note(
        total, comparison_mode=True, dim_avgs=dim_avgs, rows=rows
    )
    pos_labels = [k for k, _ in pos_dims]
    neg_labels = [k for k, _ in neg_dims]
    conclusion_one_line = (
        str(analysis.get("conclusion_one_line") or "").strip()
        or _build_conclusion_one_line(grand_note, pos_labels, neg_labels)
    )

    groups = {str(r.get("分组") or "").strip() for r in rows if r.get("分组")}
    program_types = "、".join(sorted(g for g in groups if g)) or "人声朗读、清唱、器乐等多场景"

    return {
        "test_name": test_name or str(analysis.get("test_name") or "喇叭音效对比评测"),
        "test_device": dut,
        "ref_device": ref,
        "test_mode": mode,
        "program_count": len(rows) or int(analysis.get("program_count") or 0),
        "program_types": program_types,
        "report_date": datetime.now().strftime("%Y年%m月%d日"),
        "eval_model": str(analysis.get("eval_model") or analysis.get("web_ui_eval_model") or ""),
        "total_score_avg": total,
        "dimension_scores": dimension_scores,
        "summary": summary,
        "program_details": {
            "advantage_programs": advantage_programs,
            "weakness_programs": weakness_programs,
        },
        "optimization_suggestions": {
            "dut_priorities": dut_priorities,
            "ref_suggestions": list(_REF_OPTIMIZATION_DEFAULT),
            "retest_focus": retest,
        },
        "highlights": highlights,
        "weaknesses": weaknesses,
        "conclusion_one_line": conclusion_one_line,
        "grand_performance_note": grand_note,
        "detail_rows": rows,
    }


def build_report_payload_from_analysis_file(path: str | Path) -> dict[str, Any]:
    p = Path(path)
    data = json.loads(p.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("analysis JSON 根节点须为 object")
    return build_report_payload_from_analysis(data)


def default_speaker_report_path(
    base_dir: str | Path | None = None,
    *,
    when: datetime | None = None,
) -> Path:
    """``音效对比报告_{YYYYMMDD}.docx``"""
    from config import REPORT_DIR

    dt = when or datetime.now()
    root = Path(base_dir) if base_dir else REPORT_DIR
    root.mkdir(parents=True, exist_ok=True)
    return root / f"音效对比报告_{dt.strftime('%Y%m%d')}.docx"


def _build_conclusion_one_line(
    grand_note: str,
    pos_labels: list[str],
    neg_labels: list[str],
) -> str:
    pos_txt = "、".join(pos_labels[:2]) if pos_labels else "无明显优势维度"
    neg_txt = "、".join(neg_labels[:2]) if neg_labels else "无明显短板维度"
    return f"{grand_note}；优势集中在{pos_txt}，需优先关注{neg_txt}。"


# ---------------------------------------------------------------------------
# 表现判定（第二章表格）
# ---------------------------------------------------------------------------
def _overview_performance_judgment(
    dim: str,
    avg: float,
    *,
    is_best: bool,
    is_worst: bool,
    sole_advantage: bool,
) -> str:
    """根据分差与相对排名生成「表现判定」列文案。"""
    dim_hint: dict[str, tuple[str, str]] = {
        "声音响度": ("，外放能量与动态不足", "，响度与动态较好"),
        "人声清晰度": ("，人声表现干涩", "，人声解析较好"),
        "听感舒适度": ("，久听易疲劳", "，耐听性较好"),
        "失真与噪声": ("，失真与底噪偏高", "，失真控制更优"),
        "频响平衡": ("，高低频失衡突出", "，频段衔接较自然"),
    }
    hint_neg, hint_pos = dim_hint.get(dim, ("", ""))

    if avg >= 1.0:
        base = "测试机显著优于对比机"
        suffix = hint_pos
    elif avg >= 0.5:
        base = "测试机优于对比机"
        suffix = hint_pos
    elif avg >= 0.15:
        base = "测试机略优于对比机，差异较小"
        suffix = hint_pos
    elif avg <= -1.0:
        base = "测试机严重弱于对比机"
        suffix = hint_neg
    elif avg <= -0.5:
        base = "测试机弱于对比机"
        suffix = hint_neg
    elif avg <= -0.15:
        base = "测试机略弱于对比机，差异较小"
        suffix = hint_neg
    else:
        return "测试机与对比机基本持平，差异不显著"

    text = base + suffix
    if is_worst and avg <= -0.5:
        text += "（核心短板）"
    elif is_best and avg >= 0.35:
        text += "（唯一优势维度）" if sole_advantage else "（优势维度）"
    return text


def _normalize_payload(json_data: Mapping[str, Any]) -> dict[str, Any]:
    """合并用户传入与默认值，统一键名。"""
    data = dict(json_data)
    data.setdefault("report_date", datetime.now().strftime("%Y年%m月%d日"))
    data.setdefault("test_mode", "刺激比较法（−3～+3 分差，正值=测试机优于对比机）")
    data.setdefault("program_count", 0)
    data.setdefault("program_types", "标准评测节目")
    data.setdefault("test_device", "被测设备")
    data.setdefault("ref_device", "对比设备")

    ds = data.get("dimension_scores")
    if not isinstance(ds, dict):
        ds = {}
    dimension_scores: dict[str, float] = {}
    for k in DIMENSION_KEYS:
        v = _safe_float(ds.get(k))
        if v is not None:
            dimension_scores[k] = round(v, 2)
    data["dimension_scores"] = dimension_scores

    if data.get("total_score_avg") is None and dimension_scores:
        data["total_score_avg"] = round(
            sum(dimension_scores.values()) / len(dimension_scores), 2
        )

    pd = data.get("program_details")
    if not isinstance(pd, dict):
        pd = {}
    pd.setdefault("advantage_programs", [{"name": "无数据", "total_diff": 0, "highlight": "无数据"}])
    pd.setdefault("weakness_programs", [{"name": "无数据", "total_diff": 0, "highlight": "无数据"}])
    data["program_details"] = pd

    opt = data.get("optimization_suggestions")
    if not isinstance(opt, dict):
        opt = {}
    opt.setdefault("dut_priorities", [])
    opt.setdefault("ref_suggestions", list(_REF_OPTIMIZATION_DEFAULT))
    opt.setdefault("retest_focus", ["无数据", "无数据"])
    data["optimization_suggestions"] = opt

    data.setdefault("highlights", ["无数据"])
    data.setdefault("weaknesses", ["无数据"])
    data.setdefault("summary", "")
    data.setdefault(
        "conclusion_one_line",
        data.get("grand_performance_note") or "无数据",
    )
    return data


# ---------------------------------------------------------------------------
# python-docx 样式
# ---------------------------------------------------------------------------
def _set_run_font(run: Any, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_style(
    paragraph: Any,
    *,
    font_name: str = "宋体",
    size_pt: float = 12,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    line_spacing: float = 1.5,
) -> None:
    if align is not None:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = Pt(6)
    for run in paragraph.runs:
        _set_run_font(run, font_name, size_pt, bold=bold)


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    size = 16 if level == 1 else 14
    run = p.add_run(text)
    _set_run_font(run, "黑体", size, bold=True)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def _add_body(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, "宋体", 12, bold=bold)
    _set_paragraph_style(p)


def _set_cell_border(cell: Any, **kwargs: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        tag = f"w:{edge}"
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), kwargs[edge])
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(element)


def _shade_cell(cell: Any, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _write_table_cell(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    center: bool = True,
    header: bool = False,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, "黑体" if header else "宋体", 11 if header else 10.5, bold=bold)
    if header:
        _shade_cell(cell, "DBEAFE")


def _add_data_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    col_widths_cm: Sequence[float] | None = None,
) -> Any:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for j, h in enumerate(headers):
        _write_table_cell(tbl.rows[0].cells[j], h, bold=True, header=True)
        _set_cell_border(tbl.rows[0].cells[j], top="000000", bottom="000000", left="000000", right="000000")
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            _write_table_cell(tbl.rows[i].cells[j], str(val), bold="总" in str(row[0]))
            _set_cell_border(tbl.rows[i].cells[j], top="000000", bottom="000000", left="000000", right="000000")
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return tbl


# ---------------------------------------------------------------------------
# Word 六章生成
# ---------------------------------------------------------------------------
def generate_audio_test_report(json_data: dict, output_path: str) -> None:
    """
    从规范 dict 生成《喇叭测试报告》Word 文档。

    参见 ``build_report_payload_from_analysis`` 了解如何从流水线 JSON 构造 ``json_data``。
    """
    data = _normalize_payload(json_data)
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sect = doc.sections[0]
    sect.page_height, sect.page_width = Cm(29.7), Cm(21.0)
    sect.left_margin = sect.right_margin = Cm(2.5)
    sect.top_margin = sect.bottom_margin = Cm(2.5)

    title = doc.add_heading("喇叭测试报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        _set_run_font(run, "黑体", 18, bold=True)

    # ----- 1. 测试基本信息 -----
    _add_heading(doc, "1. 测试基本信息", 1)
    n_prog = int(data.get("program_count") or 0)
    _add_body(doc, f"测试设备：{data['test_device']} vs {data['ref_device']}")
    _add_body(doc, f"测试模式：{data['test_mode']}")
    _add_body(
        doc,
        f"测试节目：共 {n_prog} 个标准节目（含 {data.get('program_types', '多场景')}）",
    )
    em = str(data.get("eval_model") or "").strip()
    if em:
        _add_body(doc, f"评测模型：{em}")
    _add_body(doc, f"报告日期：{data['report_date']}")

    # ----- 2. 核心数据概览 -----
    _add_heading(doc, "2. 核心数据概览", 1)
    dim_scores: dict[str, float] = data["dimension_scores"]
    ranked = sorted(
        [(k, dim_scores[k]) for k in DIMENSION_KEYS if k in dim_scores],
        key=lambda t: t[1],
    )
    best_k = max(ranked, key=lambda t: t[1])[0] if ranked else None
    worst_k = min(ranked, key=lambda t: t[1])[0] if ranked else None
    pos_count = sum(1 for _, v in ranked if v >= 0.35)
    sole_adv = pos_count == 1

    overview_rows: list[list[str]] = []
    for k in DIMENSION_KEYS:
        if k not in dim_scores:
            overview_rows.append([k, "无数据", "无数据"])
            continue
        avg = dim_scores[k]
        overview_rows.append(
            [
                k,
                f"{avg:+.2f}",
                _overview_performance_judgment(
                    k,
                    avg,
                    is_best=(k == best_k),
                    is_worst=(k == worst_k),
                    sole_advantage=sole_adv,
                ),
            ]
        )
    total_avg = data.get("total_score_avg")
    total_cell = f"{float(total_avg):+.2f}" if total_avg is not None else "无数据"
    grand_note = str(
        data.get("grand_performance_note")
        or grand_conclusion_note(
            _safe_float(total_avg),
            comparison_mode=True,
            dim_avgs=dim_scores,
            rows=[],
        )
    )
    overview_rows.append(["五维总均分", total_cell, grand_note])

    _add_data_table(
        doc,
        ["维度", "全节目平均分差", "表现判定"],
        overview_rows,
        col_widths_cm=(4.0, 4.0, 10.5),
    )

    # ----- 3. 节目级表现分析 -----
    _add_heading(doc, "3. 节目级表现分析", 1)
    pd = data["program_details"]
    _add_heading(doc, "3.1 优势场景（分差合计偏高）", 2)
    for item in pd.get("advantage_programs") or []:
        name = item.get("name", "无数据")
        diff = item.get("total_diff", "—")
        hl = item.get("highlight", "无数据")
        try:
            diff_s = f"{float(diff):+.2f}"
        except (TypeError, ValueError):
            diff_s = str(diff)
        _add_body(doc, f"· {name}（五维合计 {diff_s}）：{hl}")
    _add_heading(doc, "3.2 短板场景（分差合计偏低）", 2)
    for item in pd.get("weakness_programs") or []:
        name = item.get("name", "无数据")
        diff = item.get("total_diff", "—")
        hl = item.get("highlight", "无数据")
        try:
            diff_s = f"{float(diff):+.2f}"
        except (TypeError, ValueError):
            diff_s = str(diff)
        _add_body(doc, f"· {name}（五维合计 {diff_s}）：{hl}")

    # ----- 4. 综合评价（固定 Markdown 模板 → Word） -----
    _add_heading(doc, "4. 综合评价", 1)
    from markdown_report import _append_pairwise_comprehensive_evaluation_docx

    detail_rows = data.get("detail_rows") or []
    _append_pairwise_comprehensive_evaluation_docx(
        doc,
        dim_scores,
        data.get("total_score_avg"),
        detail_rows if isinstance(detail_rows, list) else [],
    )

    # ----- 5. 综合结论（摘要） -----
    _add_heading(doc, "5. 综合结论", 1)
    _add_body(doc, f"{data.get('conclusion_one_line', '无数据')}", bold=True)

    doc.save(str(out))


# ---------------------------------------------------------------------------
# 流水线集成
# ---------------------------------------------------------------------------
def generate_speaker_report_from_analysis_file(
    analysis_path: str | Path,
    output_path: str | Path | None = None,
    *,
    test_device: str = "",
    ref_device: str = "",
) -> Path:
    """读取 analysis JSON → 生成《喇叭测试报告》Word。"""
    payload = build_report_payload_from_analysis_file(analysis_path)
    if test_device:
        payload["test_device"] = test_device
    if ref_device:
        payload["ref_device"] = ref_device
    out = Path(output_path) if output_path else default_speaker_report_path()
    generate_audio_test_report(payload, str(out))
    return out


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print(
            "用法: python audio_test_report.py <analysis.json> [output.docx]\n"
            "示例: python audio_test_report.py output/analysis/analysis_xxx.json"
        )
        raise SystemExit(1)
    ap = Path(sys.argv[1])
    op = Path(sys.argv[2]) if len(sys.argv) > 2 else default_speaker_report_path()
    path = generate_speaker_report_from_analysis_file(ap, op)
    print(f"已生成: {path}")
