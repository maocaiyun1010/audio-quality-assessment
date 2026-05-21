# -*- coding: utf-8 -*-
"""
评测结果 Markdown 导出：逐条节目表 + 维度统计表 + 整体评述（便于复制到 Excel）。
"""
from __future__ import annotations

import html
import math
from collections import Counter
from typing import Any, Mapping, Sequence

# 固定五维（与 Dify / scoring 一致）
DIMENSION_KEYS: tuple[str, ...] = (
    "声音响度",
    "人声清晰度",
    "听感舒适度",
    "失真与噪声",
    "频响平衡",
)

_DIM_SHORT = {
    "声音响度": "响度",
    "人声清晰度": "人声",
    "听感舒适度": "舒适度",
    "失真与噪声": "失真与噪声",
    "频响平衡": "频响",
}


def _cell(s: Any) -> str:
    t = "" if s is None else str(s)
    t = t.replace("|", "｜").replace("\r", " ").replace("\n", " ")
    return t.strip() or "—"


def _dim_numeric(row: Mapping[str, Any], key: str) -> float | None:
    v = row.get(key)
    if v is None or v == "" or v == "—":
        return None
    if isinstance(v, (int, float)) and not (isinstance(v, float) and math.isnan(v)):
        return float(v)
    try:
        return float(str(v).strip())
    except ValueError:
        return None


def compute_dimension_statistics(
    rows: Sequence[Mapping[str, Any]],
) -> tuple[dict[str, float | None], float | None]:
    """各维度在有效行上的算术平均（列保留 2 位小数）；总平均由未四舍五入的列均值再平均后保留 2 位。"""
    raw_col: dict[str, float | None] = {}
    avgs: dict[str, float | None] = {}
    for k in DIMENSION_KEYS:
        vals = [_dim_numeric(r, k) for r in rows]
        nums = [x for x in vals if x is not None]
        if not nums:
            raw_col[k] = None
            avgs[k] = None
            continue
        m = sum(nums) / len(nums)
        raw_col[k] = m
        avgs[k] = round(m, 2)
    parts = [raw_col[k] for k in DIMENSION_KEYS if raw_col[k] is not None]
    grand: float | None = round(sum(parts) / len(parts), 2) if parts else None
    return avgs, grand


def one_line_summary_comparison(dim_avgs: Mapping[str, float | None]) -> str:
    """多设备刺激比较：基于五维平均分差的一句话（正=被测优）。"""
    pieces: list[str] = []
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            continue
        short = _DIM_SHORT[k]
        if a >= 0.51:
            pieces.append(f"{short}上被测明显优于对比")
        elif a <= -0.51:
            pieces.append(f"{short}上被测明显逊于对比")
        elif a > 0.05:
            pieces.append(f"{short}上被测略优于对比")
        elif a < -0.05:
            pieces.append(f"{short}上对比略优于被测")
    if not pieces:
        return "综合五维平均分差，被测与对比整体听感接近，未见稳定单向拉开。"
    if len(pieces) <= 4:
        return "本次评测中，" + "，".join(pieces) + "。"
    return "本次评测中，" + "，".join(pieces[:4]) + "等。"


def one_line_summary_single(grand: float | None, n_tracks: int) -> str:
    if grand is None:
        return f"共 {n_tracks} 条节目，五维可统计分不足，请检查模型是否返回完整 JSON。"
    return (
        f"共 {n_tracks} 条节目，五维主观均分（各节目五维再跨节目平均）为 {grand:.2f} 分（1–10 分制），"
        "具体各维度差异见上表。"
    )


_VERDICT_SET = frozenset({"优于", "相当", "劣于"})


def _verdict_from_grand_pairwise(grand: float | None) -> str:
    if grand is None:
        return "相当"
    if grand > 0.08:
        return "优于"
    if grand < -0.08:
        return "劣于"
    return "相当"


def _pairwise_verdict_merged(
    grand: float | None,
    rows: Sequence[Mapping[str, Any]],
) -> str:
    votes = [str(r.get("综合结论") or "").strip() for r in rows]
    votes = [v for v in votes if v in _VERDICT_SET]
    gver = _verdict_from_grand_pairwise(grand)
    if not votes:
        return gver
    mode = Counter(votes).most_common(1)[0][0]
    if grand is not None and abs(float(grand)) >= 0.2:
        return gver
    return mode


def _clip_cn_text(s: str, lo: int, hi: int) -> str:
    s = (s or "").strip().replace("\n", "")
    if len(s) > hi:
        cut = s[:hi]
        for punct in ("。", "；", "，"):
            i = cut.rfind(punct)
            if i >= max(40, lo - 30):
                return cut[: i + 1]
        return cut
    pad = "建议结合典型节目与不同音量档复听复核，以验证上述统计趋势。"
    while len(s) < lo and len(s) + len(pad) + 1 <= hi:
        s = s + pad
    return s[:hi]


def _track_dim_sum(row: Mapping[str, Any]) -> float:
    return sum(_dim_numeric(row, k) or 0.0 for k in DIMENSION_KEYS)


def _dim_trend_phrase(k: str, avg: float) -> str:
    short = _DIM_SHORT.get(k, k)
    if avg > 0.05:
        return f"{k}（{short}）平均分差 {avg:+.2f}，测试机优于对比机"
    if avg < -0.05:
        return f"{k}（{short}）平均分差 {avg:+.2f}，测试机逊于对比机"
    return f"{k}（{short}）平均分差 {avg:+.2f}，两机相当"


def _pairwise_diff_tier(avg: float) -> str:
    if avg >= 1.0:
        return "strong_pos"
    if avg >= 0.35:
        return "mild_pos"
    if avg <= -1.0:
        return "strong_neg"
    if avg <= -0.35:
        return "mild_neg"
    return "flat"


_DIM_PAIRWISE_CONCLUSION: dict[str, dict[str, str]] = {
    "声音响度": {
        "strong_pos": "响度饱满，显著优于对比机",
        "mild_pos": "略优于对比机",
        "flat": "两机持平",
        "mild_neg": "略弱于对比机",
        "strong_neg": "显著弱于对比机",
    },
    "人声清晰度": {
        "strong_pos": "人声解析明显更好",
        "mild_pos": "人声略优",
        "flat": "两机相当",
        "mild_neg": "人声解析略逊",
        "strong_neg": "存在明显短板",
    },
    "听感舒适度": {
        "strong_pos": "舒适度明显更佳",
        "mild_pos": "听感更舒适",
        "flat": "表现稳定",
        "mild_neg": "刺激性略强、耐听性偏弱",
        "strong_neg": "为主要短板维度",
    },
    "失真与噪声": {
        "strong_pos": "失真控制明显更优",
        "mild_pos": "失真控制更优",
        "flat": "两机表现相当",
        "mild_neg": "失真与底噪略逊",
        "strong_neg": "失真与噪声控制偏弱",
    },
    "频响平衡": {
        "strong_pos": "衔接自然顺滑",
        "mild_pos": "频段衔接较自然",
        "flat": "频响均衡相当",
        "mild_neg": "频段衔接略逊",
        "strong_neg": "高低频失衡问题突出",
    },
}

_DIM_PAIRWISE_ROLE: dict[str, tuple[str, str]] = {
    "人声清晰度": ("为核心优势维度", "存在明显短板"),
    "听感舒适度": ("舒适度表现突出", "为主要短板维度"),
    "声音响度": ("响度为相对优势", "响度为相对短板"),
    "失真与噪声": ("失真控制为相对优势", "失真控制为相对短板"),
    "频响平衡": ("频响衔接为相对优势", "频响衔接为相对短板"),
}


def _single_score_tier(avg: float) -> str:
    if avg >= 8.5:
        return "strong_high"
    if avg >= 7.0:
        return "high"
    if avg >= 5.5:
        return "mid"
    if avg >= 4.0:
        return "low"
    return "very_low"


_DIM_SINGLE_CONCLUSION: dict[str, dict[str, str]] = {
    "声音响度": {
        "strong_high": "外放响度与动态表现优秀",
        "high": "响度表现良好",
        "mid": "响度中等，仍有提升空间",
        "low": "响度偏弱，建议加强低频与动态",
        "very_low": "响度明显不足",
    },
    "人声清晰度": {
        "strong_high": "人声通透，为核心优势维度",
        "high": "人声清晰度较好",
        "mid": "人声可懂度中等",
        "low": "人声略闷，存在明显短板",
        "very_low": "人声解析明显不足",
    },
    "听感舒适度": {
        "strong_high": "久听舒适，表现稳定",
        "high": "听感舒适度良好",
        "mid": "舒适度中等",
        "low": "高频刺激或低频轰感需优化",
        "very_low": "为主要短板维度",
    },
    "失真与噪声": {
        "strong_high": "失真与底噪控制优秀",
        "high": "失真控制良好",
        "mid": "失真控制一般",
        "low": "大音量失真或底噪需改善",
        "very_low": "失真与噪声问题突出",
    },
    "频响平衡": {
        "strong_high": "全频均衡，衔接自然顺滑",
        "high": "频响较均衡",
        "mid": "频响基本可用",
        "low": "高低频略有失衡",
        "very_low": "高低频失衡问题突出",
    },
}


def build_dim_conclusion_notes(
    dim_avgs: Mapping[str, float | None],
    *,
    comparison_mode: bool,
) -> dict[str, str]:
    """按维度生成五维评分结果表「结论说明」列文案。"""
    ranked: list[tuple[str, float]] = []
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            continue
        ranked.append((k, float(a)))
    best_k = max(ranked, key=lambda t: t[1])[0] if ranked else None
    worst_k = min(ranked, key=lambda t: t[1])[0] if ranked else None
    notes: dict[str, str] = {}
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            notes[k] = "—"
            continue
        avg = float(a)
        if comparison_mode:
            tier = _pairwise_diff_tier(avg)
            if k == best_k and avg >= 0.35 and k in _DIM_PAIRWISE_ROLE:
                notes[k] = _DIM_PAIRWISE_ROLE[k][0]
            elif k == worst_k and avg <= -0.35 and k in _DIM_PAIRWISE_ROLE:
                notes[k] = _DIM_PAIRWISE_ROLE[k][1]
            else:
                notes[k] = _DIM_PAIRWISE_CONCLUSION.get(k, {}).get(
                    tier, _dim_trend_phrase(k, avg).split("，", 1)[-1]
                )
        else:
            tier = _single_score_tier(avg)
            if k == best_k and avg >= 7.5 and tier in ("strong_high", "high"):
                notes[k] = _DIM_SINGLE_CONCLUSION.get(k, {}).get(
                    "strong_high", "相对突出的维度"
                )
            elif k == worst_k and avg < 6.0:
                notes[k] = _DIM_SINGLE_CONCLUSION.get(k, {}).get(
                    "very_low",
                    _DIM_SINGLE_CONCLUSION.get(k, {}).get("low", "建议优先复核"),
                )
            else:
                notes[k] = _DIM_SINGLE_CONCLUSION.get(k, {}).get(tier, "—")
    return notes


def grand_conclusion_note(
    grand: float | None,
    *,
    comparison_mode: bool,
    dim_avgs: Mapping[str, float | None],
    rows: Sequence[Mapping[str, Any]],
) -> str:
    if grand is None:
        return "—"
    if comparison_mode:
        verdict = _pairwise_verdict_merged(grand, rows)
        g = float(grand)
        if verdict == "优于":
            if g >= 0.8:
                return "测试机整体明显优于对比机"
            return "测试机整体略优于对比机"
        if verdict == "劣于":
            if g <= -0.8:
                return "测试机整体明显劣于对比机"
            return "测试机整体略劣于对比机"
        return "测试机与对比机整体听感相当"
    g = float(grand)
    if g >= 8.0:
        return "五维综合表现优秀"
    if g >= 6.5:
        return "五维综合表现良好"
    if g >= 5.0:
        return "五维综合表现中等，仍有优化空间"
    return "五维综合偏弱，建议优先调音复核"


# 《综合评价》固定 Markdown 模板（刺激比较）；判定词仅允许五种表述。
_TEMPLATE_VERDICT_WORDS = ("显著弱于", "略弱于", "持平", "略优于", "显著优于")

_TEMPLATE_DIM_CORE_ISSUE: dict[str, str] = {
    "声音响度": "外放响度与动态余量不足",
    "人声清晰度": "人声解析偏弱、咬字不够清晰",
    "听感舒适度": "耐听性不足、易疲劳",
    "失真与噪声": "失真与底噪控制偏弱",
    "频响平衡": "高低频失衡、频段衔接不自然",
}

_TEMPLATE_DIM_HIGHLIGHT: dict[str, str] = {
    "声音响度": "外放响度与能量感较好",
    "人声清晰度": "人声清晰度较好、可懂度稳定",
    "听感舒适度": "听感舒适度较好、耐听性稳定",
    "失真与噪声": "失真控制稳定、底噪较低",
    "频响平衡": "频响衔接较自然、全频均衡",
}

_TEMPLATE_DIM_PROBLEM_LINE: dict[str, str] = {
    "声音响度": "外放响度与饱满度不足",
    "人声清晰度": "人声解析与结像偏弱",
    "听感舒适度": "长时间聆听舒适度不足",
    "失真与噪声": "大音量失真或底噪偏高",
    "频响平衡": "高低频分配不均、过渡不够顺滑",
}

_TEMPLATE_REPORT_FOOTER = (
    "本报告基于AI音频评测系统生成，所有数据均来自标准测试音源的双设备对比录音。"
    "调音建议为听感分析的工程化指导，实际调整需结合硬件特性与主观复听验证。"
)


def template_dim_verdict(avg: float) -> str:
    """表现判定：显著弱于/略弱于/持平/略优于/显著优于 + 对比机。"""
    if avg >= 1.0:
        return "显著优于对比机"
    if avg >= 0.35:
        return "略优于对比机"
    if avg <= -1.0:
        return "显著弱于对比机"
    if avg <= -0.35:
        return "略弱于对比机"
    return "持平"


def _row_program_label(row: Mapping[str, Any]) -> str:
    for key in ("音源名称", "节目"):
        s = str(row.get(key) or "").strip()
        if s and s != "—":
            return s
    return "无数据"


def _weakness_program_names(rows: Sequence[Mapping[str, Any]], limit: int = 3) -> str:
    ranked = sorted(rows, key=_track_dim_sum)
    names: list[str] = []
    for r in ranked:
        if _track_dim_sum(r) >= -0.05:
            continue
        name = _row_program_label(r)
        if name != "无数据" and name not in names:
            names.append(name)
        if len(names) >= limit:
            break
    return "、".join(names) if names else "无数据"


def _retest_bullets_from_programs(program_names: Sequence[str]) -> list[str]:
    joined = "、".join(program_names)
    bullets: list[str] = []
    if any(k in joined for k in ("诵读", "赤壁", "朗读", "语声")):
        bullets.append(
            "人声朗读类节目，验证约 800 Hz–3 kHz 清晰度与约 3–6 kHz 齿音窗舒适度平衡"
        )
    if any(k in joined for k in ("船歌", "低音", "月亮", "月光")):
        bullets.append(
            "低音测试曲目，验证约 60–200 Hz 低频下潜、厚度与响度一致性"
        )
    if len(bullets) < 2:
        bullets.append(
            "标准音量与大音量两档连续播放，复核响度一致性与失真控制"
        )
    if len(bullets) < 2:
        bullets.append("典型优势与短板节目各复测一轮，确认分差趋势可重复")
    return bullets[:2]


def pairwise_dimension_score_table_md(
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]] | None = None,
) -> list[str]:
    """
    五维评分结果表 Markdown 行（与 Web「五维评分结果表」同源）。

    列：维度｜全节目平均分差（测试机相对对比机）｜结论说明；含总均分行与表后说明由调用方追加。
    """
    rows = list(rows or ())
    dim_notes = build_dim_conclusion_notes(dim_avgs, comparison_mode=True)
    grand_note = grand_conclusion_note(
        grand, comparison_mode=True, dim_avgs=dim_avgs, rows=rows
    )
    lines = [
        "| 维度 | 全节目平均分差（测试机相对对比机） | 结论说明 |",
        "| :--- | :--- | :--- |",
    ]
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        cell = f"{float(a):+.2f}" if a is not None else "—"
        lines.append(f"| {_cell(k)} | {cell} | {_cell(dim_notes.get(k, '—'))} |")
    gcell = f"{grand:+.2f}" if grand is not None else "—"
    lines.append(
        f"| **所有维度总平均分差** | **{gcell}** | **{_cell(grand_note)}** |"
    )
    return lines


_PAIRWISE_SCORE_TABLE_FOOTNOTE = (
    "> 分差算法：各维度先在所有节目中求算术平均，再对五个维度平均数求总平均；"
    "刺激比较模式下正值表示测试机优于对比机。"
)


def _append_pairwise_dimension_score_table_docx(
    doc: Any,
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]],
) -> None:
    """Word：五维评分结果表（三列，含总均分与结论说明）。"""
    rows = list(rows or ())
    dim_notes = build_dim_conclusion_notes(dim_avgs, comparison_mode=True)
    grand_note = grand_conclusion_note(
        grand, comparison_mode=True, dim_avgs=dim_avgs, rows=rows
    )
    tbl = doc.add_table(rows=1 + len(DIMENSION_KEYS) + 1, cols=3)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    hdr = tbl.rows[0].cells
    hdr[0].text = "维度"
    hdr[1].text = "全节目平均分差（测试机相对对比机）"
    hdr[2].text = "结论说明"
    try:
        for c in hdr:
            _set_docx_cell_fill(c, "DBEAFE")
    except Exception:
        pass
    for i, k in enumerate(DIMENSION_KEYS, start=1):
        cells = tbl.rows[i].cells
        cells[0].text = k
        a = dim_avgs.get(k)
        cells[1].text = f"{float(a):+.2f}" if a is not None else "—"
        cells[2].text = dim_notes.get(k, "—")
    last = tbl.rows[len(DIMENSION_KEYS) + 1].cells
    last[0].text = "所有维度总平均分差"
    last[1].text = f"{grand:+.2f}" if grand is not None else "—"
    last[2].text = grand_note
    try:
        for c in last:
            _set_docx_cell_fill(c, "EFF6FF")
    except Exception:
        pass
    doc.add_paragraph()
    doc.add_paragraph(
        "分差算法：各维度先在所有节目中求算术平均，再对五个维度平均数求总平均；"
        "刺激比较模式下正值表示测试机优于对比机。"
    )


def render_pairwise_comprehensive_evaluation_md(
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]] | None = None,
) -> str:
    """
    刺激比较模式《综合评价》固定 Markdown（结构不可变，仅填充数据）。

    供 Web UI 第六章、Word/Markdown 报告与 ``_pairwise_comprehensive_text`` 共用。
  「一、分维度表现」内嵌原五维评分结果表（三列），不再单独重复该表。
    """
    rows = list(rows or ())
    n_prog = len(rows)
    g = f"{grand:+.2f}" if grand is not None else "—"
    verdict = _pairwise_verdict_merged(grand, rows)

    weak_names = _weakness_program_names(rows)
    weak_name_list = [x.strip() for x in weak_names.split("、") if x.strip() and x != "无数据"]

    neg_dims = _pick_dims_sorted(dim_avgs, positive=False, limit=5)
    core_issues: list[str] = []
    for k, _ in neg_dims[:3]:
        core_issues.append(_TEMPLATE_DIM_CORE_ISSUE.get(k, f"{k}维度偏弱"))
    if not core_issues:
        core_issues = ["各维度整体接近持平，差异主要体现在节目间一致性"]
    core_issue_text = "；".join(core_issues)

    pos_dims = _pick_dims_sorted(dim_avgs, positive=True, limit=2)
    highlights: list[str] = []
    for k, v in pos_dims:
        highlights.append(
            f"{_TEMPLATE_DIM_HIGHLIGHT.get(k, k)}（平均分差 {v:+.2f}）"
        )
    if not highlights:
        flat = [
            k
            for k in DIMENSION_KEYS
            if dim_avgs.get(k) is not None and abs(float(dim_avgs[k])) <= 0.35
        ]
        if flat:
            highlights.append(f"{'、'.join(flat[:2])}等维度与对比机基本持平")
        else:
            highlights.append("无显著优势维度")

    problems: list[str] = []
    for k, v in neg_dims[:3]:
        problems.append(
            f"{k}（{v:+.2f}）：{_TEMPLATE_DIM_PROBLEM_LINE.get(k, '表现偏弱')}"
        )
    while len(problems) < 3:
        if neg_dims and len(problems) < len(neg_dims):
            break
        problems.append("无显著短板维度")
        if len(problems) >= 3:
            break

    pri_dims = sorted(neg_dims, key=lambda t: abs(t[1]), reverse=True)[:3]
    if not pri_dims:
        pri_dims = [(DIMENSION_KEYS[0], 0.0)]

    opt_blocks: list[str] = []
    for i, (k, v) in enumerate(pri_dims[:3], start=1):
        sug = _DIM_TUNING_DUT_WORSE.get(k, f"针对{k}频段做EQ与腔体优化。")
        opt_blocks.append(
            f"{i}.  **{k}（分差：{v:+.2f}）**\n    - 建议：{sug}"
        )
    while len(opt_blocks) < 3:
        k = DIMENSION_KEYS[len(opt_blocks) % len(DIMENSION_KEYS)]
        opt_blocks.append(
            f"{len(opt_blocks) + 1}.  **{k}（分差：—）**\n"
            f"    - 建议：{_DIM_TUNING_DUT_WORSE.get(k, '结合主观复听做小范围验证。')}"
        )

    retest = _retest_bullets_from_programs(weak_name_list or [weak_names])

    lines: list[str] = [
        f"本次评测共 {n_prog} 个标准节目，五维分差全节目总均分为 {g}。",
        f"综合判定：测试机相对对比机整体：{verdict}",
        "",
        "## 一、分维度表现（全节目平均分差）",
        *pairwise_dimension_score_table_md(dim_avgs, grand, rows),
        "",
        _PAIRWISE_SCORE_TABLE_FOOTNOTE,
        "",
        "## 二、节目级表现分析",
        "### 短板节目（分差合计偏低）",
        f"- 代表节目：{weak_names}",
        f"- 核心问题：{core_issue_text}",
        "",
        "## 三、核心听感特征",
        "### 优势亮点",
    ]
    for h in highlights[:2]:
        lines.append(f"- {h}")
    lines.extend(["", "### 主要问题"])
    for p in problems[:3]:
        lines.append(f"- {p}")
    lines.extend(
        [
            "",
            "## 四、调音与复测建议",
            "### 优先优化方向（按优先级排序）",
            *opt_blocks,
            "",
            "### 复测验证重点",
        ]
    )
    for r in retest:
        lines.append(f"- {r}")
    lines.extend(["", "## 五、报告说明", _TEMPLATE_REPORT_FOOTER])
    return "\n".join(lines)


def _append_pairwise_comprehensive_evaluation_docx(
    doc: Any,
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]],
) -> None:
    """Word：按固定综合评价结构写入（与 Markdown 模板一致）。"""
    rows = list(rows or ())
    n_prog = len(rows)
    g = f"{grand:+.2f}" if grand is not None else "—"
    verdict = _pairwise_verdict_merged(grand, rows)

    doc.add_paragraph(
        f"本次评测共 {n_prog} 个标准节目，五维分差全节目总均分为 {g}。"
    )
    doc.add_paragraph(f"综合判定：测试机相对对比机整体：{verdict}")

    doc.add_heading("一、分维度表现（全节目平均分差）", level=3)
    _append_pairwise_dimension_score_table_docx(doc, dim_avgs, grand, rows)

    doc.add_heading("二、节目级表现分析", level=3)
    doc.add_heading("短板节目（分差合计偏低）", level=4)
    weak_names = _weakness_program_names(rows)
    neg_dims = _pick_dims_sorted(dim_avgs, positive=False, limit=3)
    core_issues = "；".join(
        _TEMPLATE_DIM_CORE_ISSUE.get(k, f"{k}维度偏弱") for k, _ in neg_dims
    ) or "各维度整体接近持平"
    doc.add_paragraph(f"代表节目：{weak_names}", style="List Bullet")
    doc.add_paragraph(f"核心问题：{core_issues}", style="List Bullet")

    doc.add_heading("三、核心听感特征", level=3)
    doc.add_heading("优势亮点", level=4)
    pos_dims = _pick_dims_sorted(dim_avgs, positive=True, limit=2)
    if pos_dims:
        for k, v in pos_dims:
            doc.add_paragraph(
                f"{_TEMPLATE_DIM_HIGHLIGHT.get(k, k)}（平均分差 {v:+.2f}）",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("无显著优势维度", style="List Bullet")

    doc.add_heading("主要问题", level=4)
    for k, v in neg_dims[:3]:
        doc.add_paragraph(
            f"{k}（{v:+.2f}）：{_TEMPLATE_DIM_PROBLEM_LINE.get(k, '表现偏弱')}",
            style="List Bullet",
        )
    for _ in range(max(0, 3 - len(neg_dims))):
        doc.add_paragraph("无显著短板维度", style="List Bullet")

    doc.add_heading("四、调音与复测建议", level=3)
    doc.add_heading("优先优化方向（按优先级排序）", level=4)
    pri_dims = sorted(neg_dims, key=lambda t: abs(t[1]), reverse=True)[:3]
    for i, (k, v) in enumerate(pri_dims or [(DIMENSION_KEYS[0], 0.0)], start=1):
        doc.add_paragraph(
            f"{i}. {k}（分差：{v:+.2f}）：{_DIM_TUNING_DUT_WORSE.get(k, '')}",
            style="List Number",
        )

    doc.add_heading("复测验证重点", level=4)
    weak_name_list = [x.strip() for x in weak_names.split("、") if x.strip()]
    for r in _retest_bullets_from_programs(weak_name_list):
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("五、报告说明", level=3)
    doc.add_paragraph(_TEMPLATE_REPORT_FOOTER)


def _pairwise_comprehensive_paragraphs(
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]],
) -> list[str]:
    """刺激比较模式：多段综合评价（五维逐项 + 逐节目统计 + 典型节目）。"""
    n_prog = len(rows)
    g = f"{grand:.2f}" if grand is not None else "—"
    verdict = _pairwise_verdict_merged(grand, rows)

    paras: list[str] = [
        (
            f"本次共 {n_prog} 个节目参与统计。五维分差全节目总均分为 {g} "
            f"（正值表示测试机优于对比机，负值表示测试机整体偏弱）。"
            f"综合判定：测试机相对对比机整体「{verdict}」。"
        )
    ]

    dim_bits: list[str] = []
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            continue
        dim_bits.append(_dim_trend_phrase(k, float(a)))
    if dim_bits:
        paras.append("分维度全节目平均：" + "；".join(dim_bits) + "。")

    pos = _pick_dims_sorted(dim_avgs, positive=True, limit=5)
    neg = _pick_dims_sorted(dim_avgs, positive=False, limit=5)
    flat = [
        k
        for k in DIMENSION_KEYS
        if dim_avgs.get(k) is not None and abs(float(dim_avgs[k])) <= 0.05
    ]

    if pos:
        plab = "、".join(f"{_DIM_SHORT[k]}（{v:+.2f}）" for k, v in pos)
        paras.append(f"测试机相对领先的维度：{plab}。")
    if neg:
        nlab = "、".join(f"{_DIM_SHORT[k]}（{v:+.2f}）" for k, v in neg)
        paras.append(
            f"测试机相对落后的维度：{nlab}；调音与复测建议优先针对上述维度及对应频段。"
        )
    if flat:
        paras.append(f"接近持平、差异不显著的维度：{'、'.join(flat)}。")
    if not pos and not neg and dim_bits:
        paras.append(
            "各维度平均分差均未明显偏离 0，整体听感接近，差异主要体现在节目间一致性而非单一维度单向拉开。"
        )

    vcount = Counter(
        str(r.get("综合结论") or "").strip()
        for r in rows
        if str(r.get("综合结论") or "").strip() in _VERDICT_SET
    )
    if vcount:
        vparts = [f"「{k}」{v} 条" for k, v in vcount.most_common()]
        paras.append("逐节目「综合结论」统计：" + "，".join(vparts) + "。")

    if rows:
        ranked = sorted(rows, key=_track_dim_sum, reverse=True)
        best = [
            str(r.get("音源名称") or "—").strip()
            for r in ranked[:3]
            if _track_dim_sum(r) > 0.05
        ]
        worst = [
            str(r.get("音源名称") or "—").strip()
            for r in sorted(rows, key=_track_dim_sum)[:3]
            if _track_dim_sum(r) < -0.05
        ]
        if best:
            paras.append(
                "测试机优势较明显的节目（五维分差合计偏高）："
                + "、".join(best)
                + "。"
            )
        if worst:
            paras.append(
                "测试机短板较突出的节目（五维分差合计偏低）："
                + "、".join(worst)
                + "。"
            )

    tips: list[str] = []
    seen: set[str] = set()
    for r in rows:
        for key in ("专业点评", "对比总结", "综合评价"):
            t = str(r.get(key) or "").strip()
            if not t or t == "—" or t in seen or len(t) < 12:
                continue
            seen.add(t)
            tips.append(t[:80] + ("…" if len(t) > 80 else ""))
            if len(tips) >= 2:
                break
        if len(tips) >= 2:
            break
    if tips:
        paras.append("模型逐条听感摘要（节选）：" + "；".join(tips) + "。")

    paras.append(
        "建议对照上文五维表与各节目明细，在标准音量与大音量两档对典型优势/短板节目复听，"
        "确认统计趋势与主观听感一致后再做调音决策。"
    )
    return paras


# 平均分差为负：测试机该维偏弱 → 测试机侧调音方向。
_DIM_TUNING_DUT_WORSE: dict[str, str] = {
    "声音响度": "在约60–200Hz段适度提升低频下潜与箱体辐射效率，并控制大动态下2–4kHz人声临场段的压缩感，以抬升主观响度与厚度。",
    "人声清晰度": "在约800Hz–3kHz人声主体段提升清晰度与口腔结像，并收敛约3–6kHz齿音与过量细节峰，改善可懂度与耐听度。",
    "听感舒适度": "抑制约3–8kHz尖刺与齿音刺激，同时减轻约80–250Hz闷轰与腔体共振，降低长时间聆听的疲劳与压迫感。",
    "失真与噪声": "在常用与大音量工况下压低约100Hz–1kHz谐波与调制失真，并降低静态底噪与结构共振引起的附加噪声。",
    "频响平衡": "拉齐约200Hz–2kHz中低频与中频衔接，避免4–8kHz孤立抬升或2–5kHz凹陷，使全频过渡顺滑、声像稳定。",
}

# 平均分差为正：对比机该维偏弱 → 对比机侧调音方向。
_DIM_TUNING_REF_WORSE: dict[str, str] = {
    "声音响度": "补强约60–200Hz低频下潜与前障板辐射效率，并检视约2–4kHz人声临场段是否偏薄导致主观响度不足。",
    "人声清晰度": "提升约800Hz–3kHz人声主体清晰度，适度增强约3–5kHz唇齿细节的自然度，避免仅靠整体音量补偿带来的粗糙感。",
    "听感舒适度": "抑制约3–7kHz尖刺与齿音过量，并减轻约80–200Hz轰头与低频堆积，提高多素材下的耐听一致性。",
    "失真与噪声": "针对大动态下约200Hz–1.5kHz谐波堆积与底噪抬升，优化换能与腔体耦合，降低可闻失真与噪声掩蔽。",
    "频响平衡": "矫正约200Hz–1kHz过量或不足，并理顺约4–9kHz高频延伸与空气感，使高中低衔接更自然。",
}


def _pick_dims_sorted(
    dim_avgs: Mapping[str, float | None],
    *,
    positive: bool,
    limit: int,
) -> list[tuple[str, float]]:
    out: list[tuple[str, float]] = []
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            continue
        if positive and a > 0.05:
            out.append((k, float(a)))
        if not positive and a < -0.05:
            out.append((k, float(a)))
    out.sort(key=lambda t: (-t[1] if positive else t[1]))
    return out[:limit]


def _bullets_dut_all_positive_or_flat(dim_avgs: Mapping[str, float | None]) -> list[str]:
    """测试机五维平均无逆差时：给出巩固优势与一致性的可执行建议。"""
    ranked = _pick_dims_sorted(dim_avgs, positive=True, limit=3)
    lead = "、".join(f"{_DIM_SHORT[k]}" for k, _ in ranked) if ranked else "领先维度"
    return [
        (
            f"- 在已占优的{lead}上保持现有声学取向，仅在约2–5kHz人声齿音窗做±0.5dB量级细调，"
            "避免为追求解析而引入新的刺耳感。"
        ),
        "- 在大音量档复核约60–150Hz动态余量与约3–6kHz峰值的曲间一致性，降低不同节目下的听感波动。",
        "- 与对比机保持同一摆位与采集链路做一次复测，确认领先幅度在不同素材上可重复、非偶然摆位引入。",
    ]


def _bullets_for_pairwise(
    dim_avgs: Mapping[str, float | None],
    *,
    for_dut: bool,
) -> list[str]:
    if for_dut:
        dims = _pick_dims_sorted(dim_avgs, positive=False, limit=5)
        if not dims:
            return _bullets_dut_all_positive_or_flat(dim_avgs)
        return [
            f"- {_DIM_TUNING_DUT_WORSE.get(k, k)}（对应维度：{k}，节目平均分差 {v:+.2f}）"
            for k, v in dims[:3]
        ]
    dims = _pick_dims_sorted(dim_avgs, positive=True, limit=5)
    return [f"- {_DIM_TUNING_REF_WORSE.get(k, k)}（对应维度：{k}，节目平均分差 {v:+.2f}）" for k, v in dims[:3]]


def _fallback_bullets_pairwise(dim_avgs: Mapping[str, float | None], *, for_dut: bool) -> list[str]:
    ranked: list[tuple[str, float]] = []
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            continue
        ranked.append((k, float(a)))
    ranked.sort(key=lambda t: abs(t[1]), reverse=True)
    out: list[str] = []
    for k, v in ranked[:3]:
        if for_dut and v < 0:
            out.append(
                f"- {_DIM_TUNING_DUT_WORSE.get(k, k)}（对应维度：{k}，节目平均分差 {v:+.2f}）"
            )
        if not for_dut and v > 0:
            out.append(
                f"- {_DIM_TUNING_REF_WORSE.get(k, k)}（对应维度：{k}，节目平均分差 {v:+.2f}）"
            )
    if len(out) < 2:
        tip = (
            "- 在约200Hz–4kHz人声与乐感关键窗做细粒度EQ与相位检查，结合腔体谐振抑制，保证与对端设备听感对齐。"
            if for_dut
            else "- 在约60Hz–2kHz能量分布与约3–8kHz细节窗同步优化，缩小与测试机的主观差距。"
        )
        out.append(tip)
    extras = (
        "- 在约1–4kHz人声与乐器存在感窗复核EQ与相位，并配合约60–150Hz腔体谐振抑制，降低房间边界引起的听感误判。",
        "- 在标准音量与大音量两档复测声压与谱形，核实差异来自终端本体而非单次摆位偶然因素。",
    )
    ei = 0
    while len(out) < 2 and ei < len(extras):
        out.append(extras[ei])
        ei += 1
    return out[:3]


def _pairwise_comprehensive_text(
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]] | None = None,
) -> str:
    return render_pairwise_comprehensive_evaluation_md(dim_avgs, grand, rows)


def _single_comprehensive_text(
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]] | None = None,
) -> str:
    rows = rows or ()
    g = f"{grand:.2f}" if grand is not None else "—"
    ranked: list[tuple[str, float]] = []
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            continue
        ranked.append((k, float(a)))
    ranked.sort(key=lambda t: -t[1])
    weak = sorted(ranked, key=lambda t: t[1])[:3]
    paras = [
        (
            f"单终端模式共 {len(rows)} 个节目，五维总平均 {g} 分（1–10 分制）。"
            "以下为各维度全节目平均得分及相对高低。"
        )
    ]
    if ranked:
        all_bits = [
            f"{k}（{_DIM_SHORT[k]}）{v:.2f} 分" for k, v in ranked
        ]
        paras.append("分维度平均：" + "；".join(all_bits) + "。")
    if ranked:
        top = "、".join(f"{_DIM_SHORT[k]}（{v:.2f}）" for k, v in ranked[:3])
        paras.append(f"相对突出的维度：{top}。")
    if weak and ranked and weak[0][1] < ranked[0][1] - 0.3:
        wlab = "、".join(f"{_DIM_SHORT[k]}（{v:.2f}）" for k, v in weak)
        paras.append(f"相对偏弱、建议优先复核的维度：{wlab}。")
    if rows:
        hi = sorted(rows, key=_track_dim_sum, reverse=True)[:2]
        lo = sorted(rows, key=_track_dim_sum)[:2]
        if hi:
            paras.append(
                "得分较高的节目："
                + "、".join(str(r.get("音源名称") or "—") for r in hi)
                + "。"
            )
        if lo and hi and _track_dim_sum(lo[0]) < _track_dim_sum(hi[0]) - 0.1:
            paras.append(
                "得分偏低的节目："
                + "、".join(str(r.get("音源名称") or "—") for r in lo)
                + "。"
            )
    paras.append(
        "建议结合上表各节目明细，在典型音量档复听验证后开展调音迭代。"
    )
    return "\n\n".join(paras)


def _single_tuning_bullets(dim_avgs: Mapping[str, float | None]) -> list[str]:
    ranked: list[tuple[str, float]] = []
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        if a is None:
            continue
        ranked.append((k, float(a)))
    ranked.sort(key=lambda t: t[1])
    out: list[str] = []
    for k, v in ranked[:3]:
        out.append(
            f"- {_DIM_TUNING_DUT_WORSE.get(k, k)}（对应维度：{k}，节目平均分 {v:.2f}）"
        )
    if not out:
        out.append(
            "- 在约200Hz–4kHz人声与乐感关键窗做细粒度EQ与相位检查，并结合腔体谐振抑制，提升整体主观一致性。"
        )
    return out


def _core_conclusion_highlight_html(
    *,
    comparison_mode: bool,
    grand: float | None,
    rows: Sequence[Mapping[str, Any]],
) -> str:
    """核心结论：独立卡片样式（HTML，Web 端需 unsafe_allow_html）。"""
    if comparison_mode:
        verdict = _pairwise_verdict_merged(grand, rows)
        main = f"✅ 测试机 整体音质「{verdict}」对比机"
        sub = "基于全节目五维平均的判定摘要，请结合上表与综合评价阅读。"
    else:
        gshow = f"{grand:.2f}" if grand is not None else "—"
        main = (
            f"✅ 本次单终端评测五维总平均分为 {gshow} 分（1–10 分制），详见上表各维度分解。"
        )
        sub = "基于全节目均分的结论摘要，请结合上表与综合评价阅读。"
    _m = html.escape(main)
    _s = html.escape(sub)
    return f"""<div style="background:linear-gradient(165deg,#eff6ff 0%,#dbeafe 55%,#bfdbfe 100%);border:3px solid #2563eb;border-radius:14px;padding:22px 26px;margin:20px 0 24px;text-align:center;box-shadow:0 4px 14px rgba(37,99,235,0.18);">
<p style="margin:0;font-size:1.65em;font-weight:800;color:#1e3a8a;line-height:1.35;letter-spacing:0.03em;text-shadow:0 1px 0 rgba(255,255,255,0.5);">{_m}</p>
<p style="margin:14px 0 0;font-size:0.98em;color:#475569;font-weight:500;">{_s}</p>
</div>"""


def _set_docx_cell_fill(cell: Any, fill_hex: str) -> None:
    """为 Word 表格单元格设置背景色（6 位十六进制，不含 #）。"""
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), fill_hex))
        cell._tc.get_or_add_tcPr().append(shd)
    except Exception:
        pass


def build_section_six_markdown(
    *,
    comparison_mode: bool,
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]],
) -> str:
    lines: list[str] = []
    if not comparison_mode:
        lines.append("### 五维评分结果表")
        lines.append("")
        dim_notes = build_dim_conclusion_notes(dim_avgs, comparison_mode=False)
        grand_note = grand_conclusion_note(
            grand, comparison_mode=False, dim_avgs=dim_avgs, rows=rows
        )
        lines.append("| 维度 | 全节目平均分（1–10 分制） | 结论说明 |")
        lines.append("| --- | --- | --- |")
        for k in DIMENSION_KEYS:
            a = dim_avgs.get(k)
            cell = f"{a:.2f}" if a is not None else "—"
            lines.append(f"| {_cell(k)} | {cell} | {_cell(dim_notes.get(k, '—'))} |")
        gcell = f"{grand:.2f}" if grand is not None else "—"
        lines.append(f"| **所有维度总平均分** | **{gcell}** | **{_cell(grand_note)}** |")
        lines.append("")
        lines.append(
            "> 平均分算法：各维度先在所有节目中求算术平均，再对五个维度平均数求总平均。"
        )
        lines.append("")
    lines.append("### 核心结论")
    lines.append("")
    lines.append(
        _core_conclusion_highlight_html(
            comparison_mode=comparison_mode, grand=grand, rows=rows
        )
    )
    lines.append("")
    lines.append("### 综合评价")
    lines.append("")
    if comparison_mode:
        lines.append(_pairwise_comprehensive_text(dim_avgs, grand, rows))
    else:
        lines.append(_single_comprehensive_text(dim_avgs, grand, rows))
    lines.append("")
    if not comparison_mode:
        lines.append("### 优化建议")
        lines.append("")
        lines.append("#### 终端调音优化建议")
        lines.append("")
        lines.extend(_single_tuning_bullets(dim_avgs))
        lines.append("")
    try:
        from nisqa_local import render_nisqa_appendix_markdown

        appendix = render_nisqa_appendix_markdown(list(rows))
        if appendix.strip():
            lines.append(appendix.rstrip())
    except Exception:
        pass
    return "\n".join(lines)


def append_section_six_docx(
    doc: Any,
    *,
    comparison_mode: bool,
    dim_avgs: Mapping[str, float | None],
    grand: float | None,
    rows: Sequence[Mapping[str, Any]],
) -> None:
    doc.add_heading("本次评测最终结论与结果汇总", level=1)
    if comparison_mode:
        doc.add_paragraph(
            "五维分差表已并入下方「综合评价 · 一、分维度表现」；"
            "请先阅读核心结论，再查看综合评价全文。"
        )
    else:
        doc.add_heading("五维评分结果表", level=2)
        dim_notes = build_dim_conclusion_notes(dim_avgs, comparison_mode=False)
        grand_note = grand_conclusion_note(
            grand, comparison_mode=False, dim_avgs=dim_avgs, rows=rows
        )
        tbl = doc.add_table(rows=1 + len(DIMENSION_KEYS) + 1, cols=3)
        try:
            tbl.style = "Table Grid"
        except Exception:
            pass
        hdr = tbl.rows[0].cells
        hdr[0].text = "维度"
        hdr[1].text = "全节目平均分（1–10）"
        hdr[2].text = "结论说明"
        for i, k in enumerate(DIMENSION_KEYS, start=1):
            cells = tbl.rows[i].cells
            cells[0].text = k
            a = dim_avgs.get(k)
            cells[1].text = f"{a:.2f}" if a is not None else "—"
            cells[2].text = dim_notes.get(k, "—")
        last = tbl.rows[len(DIMENSION_KEYS) + 1].cells
        last[0].text = "所有维度总平均"
        last[1].text = f"{grand:.2f}" if grand is not None else "—"
        last[2].text = grand_note
        doc.add_paragraph()

    if comparison_mode:
        try:
            from docx.shared import Pt

            p_h = doc.add_paragraph()
            r_h = p_h.add_run("核心结论")
            r_h.bold = True
            r_h.font.size = Pt(13)
        except Exception:
            doc.add_paragraph("核心结论")
    else:
        doc.add_heading("核心结论", level=2)
    if comparison_mode:
        verdict = _pairwise_verdict_merged(grand, rows)
        core_text = f"✅ 测试机 整体音质「{verdict}」对比机"
        hint_text = "基于全节目五维平均的判定摘要，请结合上表与综合评价阅读。"
    else:
        gshow = f"{grand:.2f}" if grand is not None else "—"
        core_text = (
            f"✅ 本次单终端评测五维总平均分为 {gshow} 分（1–10 分制），详见上表各维度分解。"
        )
        hint_text = "基于全节目均分的结论摘要，请结合上表与综合评价阅读。"

    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        tbl_core = doc.add_table(rows=1, cols=1)
        tbl_core.autofit = True
        try:
            tbl_core.style = "Table Grid"
        except Exception:
            pass
        c0 = tbl_core.rows[0].cells[0]
        _set_docx_cell_fill(c0, "DBEAFE")
        c0_par = c0.paragraphs[0]
        c0_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c0_par.paragraph_format.space_before = Pt(10)
        c0_par.paragraph_format.space_after = Pt(4)
        r_main = c0_par.add_run(core_text)
        r_main.bold = True
        r_main.font.size = Pt(22)
        r_main.font.color.rgb = RGBColor(23, 37, 84)
        p2 = c0.add_paragraph()
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(10)
        r_sub = p2.add_run(hint_text)
        r_sub.italic = True
        r_sub.font.size = Pt(10)
        r_sub.font.color.rgb = RGBColor(71, 85, 105)
        doc.add_paragraph()
    except Exception:
        doc.add_paragraph(core_text)
        doc.add_paragraph(f"（{hint_text}）")

    doc.add_heading("综合评价", level=2)
    if comparison_mode:
        _append_pairwise_comprehensive_evaluation_docx(doc, dim_avgs, grand, rows)
    else:
        comp_paras = _single_comprehensive_text(dim_avgs, grand, rows).split("\n\n")
        for para in comp_paras:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_heading("优化建议", level=2)
        doc.add_heading("终端调音优化建议", level=3)
        for line in _single_tuning_bullets(dim_avgs):
            t = line[2:].strip() if line.startswith("- ") else line
            doc.add_paragraph(t, style="List Bullet")
    _append_nisqa_appendix_docx(doc, rows)
    doc.add_paragraph()


def _append_nisqa_appendix_docx(doc: Any, rows: Sequence[Mapping[str, Any]]) -> None:
    """Word 终章：NISQA 客观分表（与 Markdown 附录一致）。"""
    try:
        from nisqa_local import render_nisqa_appendix_markdown

        appendix = render_nisqa_appendix_markdown(list(rows))
    except Exception:
        return
    if not appendix.strip():
        return

    doc.add_heading("NISQA 客观音质（本地）", level=2)
    for line in appendix.splitlines():
        text = line.rstrip()
        if not text:
            doc.add_paragraph()
        elif text.startswith("> "):
            doc.add_paragraph(text[2:].strip())
        elif text.startswith("| "):
            doc.add_paragraph(text)
        elif text.startswith("### "):
            doc.add_heading(text[4:].strip(), level=3)
        else:
            doc.add_paragraph(text)


def render_evaluation_markdown(
    *,
    test_name: str,
    test_device: str,
    ref_device: str,
    eval_model_name: str = "",
    comparison_mode: bool,
    cross_session: bool,
    rows: list[dict[str, Any]],
    dim_avgs: dict[str, float | None],
    grand: float | None,
    one_line: str,
) -> str:
    lines: list[str] = []
    lines.append(f"# {test_name}")
    lines.append("")
    _em = (eval_model_name or "").strip()
    if _em:
        lines.append(f"- **评测大模型**：{_cell(_em)}")
    lines.append(f"- **被测 / 主测**：{_cell(test_device)}")
    lines.append(f"- **对比 / 参考**：{_cell(ref_device)}")
    if comparison_mode:
        if cross_session:
            lines.append("- **模式**：跨会话刺激比较（五维整数 −3～+3，正为被测优）")
        else:
            lines.append("- **模式**：同会话刺激比较（五维整数 −3～+3，正为被测优）")
    else:
        lines.append("- **模式**：单终端逐条评测（五维 1–10 分）")
    lines.append("")
    lines.append("## 逐条节目测评")
    lines.append("")

    if comparison_mode:
        headers = ["节目", "分组"] + list(DIMENSION_KEYS) + ["对比总结", "专业点评"]
    else:
        headers = ["节目", "分组"] + list(DIMENSION_KEYS) + ["综合分", "专业点评"]

    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

    for r in rows:
        cells: list[str] = []
        cells.append(_cell(r.get("节目")))
        cells.append(_cell(r.get("分组")))
        for k in DIMENSION_KEYS:
            v = r.get(k)
            if v is None or v == "":
                cells.append("—")
                continue
            try:
                fv = float(v)
            except (TypeError, ValueError):
                cells.append(_cell(v))
                continue
            if comparison_mode and abs(fv - round(fv)) < 1e-9:
                cells.append(str(int(round(fv))))
            else:
                cells.append(f"{fv:.2f}")
        if comparison_mode:
            cells.append(_cell(r.get("对比总结")))
            cells.append(_cell(r.get("专业点评")))
        else:
            z = r.get("综合分")
            if z is None or z == "":
                cells.append("—")
            else:
                try:
                    cells.append(f"{float(z):.2f}")
                except (TypeError, ValueError):
                    cells.append(_cell(z))
            cells.append(_cell(r.get("专业点评")))
        lines.append("| " + " | ".join(cells) + " |")

    lines.append("")
    lines.append("## 维度统计（全节目平均）")
    lines.append("")
    lines.append("| 维度 | 平均分 |")
    lines.append("| --- | --- |")
    for k in DIMENSION_KEYS:
        a = dim_avgs.get(k)
        cell = f"{a:.2f}" if a is not None else "—"
        lines.append(f"| {_cell(k)} | {cell} |")
    gcell = f"{grand:.2f}" if grand is not None else "—"
    lines.append(f"| **所有维度总平均分** | **{gcell}** |")
    lines.append("")
    lines.append("> 总平均分算法：先对每个维度在所有节目中求算术平均，再对五个维度平均数再求平均（仅含有效维度）。")
    lines.append("")
    lines.append("## 整体评述")
    lines.append("")
    lines.append(one_line)
    lines.append("")
    lines.append("## 本次评测最终结论与结果汇总")
    lines.append("")
    lines.append(
        build_section_six_markdown(
            comparison_mode=comparison_mode,
            dim_avgs=dim_avgs,
            grand=grand,
            rows=rows,
        )
    )
    return "\n".join(lines)


def write_markdown(path: str, content: str) -> None:
    from pathlib import Path

    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(content, encoding="utf-8")
